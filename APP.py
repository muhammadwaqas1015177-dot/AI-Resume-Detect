import io
import json
import os
import re
from pathlib import Path
from typing import Any

import streamlit as st
from google import genai
from google.genai import types
from pypdf import PdfReader
from docx import Document

st.set_page_config(page_title="Resume ATS Analyzer", page_icon="📄", layout="wide")

MODEL_NAME = "gemini-2.5-flash"
MAX_RESUME_CHARS = 50000

WEIGHTS = {
    "ats_format": 20,
    "sections": 15,
    "keywords": 20,
    "experience_bullets": 20,
    "skills": 10,
    "contact": 5,
    "education": 5,
    "readability": 5,
}


def get_api_key() -> str | None:
    try:
        key = st.secrets.get("GEMINI_API_KEY")
        if key:
            return str(key).strip()
    except Exception:
        pass
    return os.getenv("GEMINI_API_KEY")


def extract_pdf(file_bytes: bytes) -> str:
    reader = PdfReader(io.BytesIO(file_bytes))
    return "\n".join((page.extract_text() or "") for page in reader.pages).strip()


def extract_docx(file_bytes: bytes) -> str:
    doc = Document(io.BytesIO(file_bytes))
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()


def extract_text(uploaded_file) -> str:
    data = uploaded_file.getvalue()
    suffix = Path(uploaded_file.name).suffix.lower()

    if suffix == ".pdf":
        text = extract_pdf(data)
    elif suffix == ".docx":
        text = extract_docx(data)
    elif suffix in {".txt", ".md"}:
        text = data.decode("utf-8", errors="replace").strip()
    else:
        raise ValueError("Unsupported file type.")

    if not text:
        raise ValueError(
            "No readable text was found. If this is a scanned/image-only PDF, "
            "upload an OCR-enabled PDF or DOCX file."
        )

    return text[:MAX_RESUME_CHARS]


def build_prompt(resume: str, job_description: str) -> str:
    jd = job_description.strip()
    job_part = (
        f"JOB DESCRIPTION:\n{jd[:30000]}"
        if jd
        else
        "JOB DESCRIPTION:\nNot provided. Evaluate general ATS readiness."
    )

    return f"""
You are an expert ATS resume reviewer and technical recruiter.

Analyze the resume below for ATS-readiness and recruiter readability.

Rules:
- Base findings only on supplied resume/job description.
- Never invent experience, employers, dates, degrees, certifications, skills,
  metrics, achievements, or keywords.
- Do not recommend keyword stuffing.
- A missing keyword should only be recommended if the candidate genuinely has
  that skill/experience.
- Focus on likely ATS risks such as unusual section headings, columns/tables,
  graphics/icons, headers/footers, unclear dates, and poor text structure.
- If a job description is provided, distinguish keywords that are present from
  keywords that are missing or weak.
- Bullet rewrites must use only facts already present in the resume.
- Return ONLY valid JSON.

Score each category from 0 to 100:
ats_format: parser-friendly structure and formatting
sections: completeness and clarity of standard resume sections
keywords: job-description alignment, or general role-relevant terminology
experience_bullets: action verbs, specificity, outcomes, metrics, concision
skills: clarity and relevance of skills
contact: clear professional contact information
education: clarity of education details
readability: grammar, consistency, concision, scanability

Return this exact JSON structure:
{{
  "summary": "string",
  "category_scores": {{
    "ats_format": 0,
    "sections": 0,
    "keywords": 0,
    "experience_bullets": 0,
    "skills": 0,
    "contact": 0,
    "education": 0,
    "readability": 0
  }},
  "category_evidence": {{
    "ats_format": "string",
    "sections": "string",
    "keywords": "string",
    "experience_bullets": "string",
    "skills": "string",
    "contact": "string",
    "education": "string",
    "readability": "string"
  }},
  "strengths": ["string"],
  "improvements": [
    {{
      "priority": "High|Medium|Low",
      "issue": "string",
      "recommendation": "string"
    }}
  ],
  "keyword_analysis": {{
    "present": ["string"],
    "missing_or_weak": ["string"],
    "notes": "string"
  }},
  "bullet_rewrites": [
    {{
      "original": "string",
      "improved": "string",
      "reason": "string"
    }}
  ],
  "ats_checklist": ["string"]
}}

RESUME:
{resume}

{job_part}
"""


def normalize_result(result: dict[str, Any]) -> dict[str, Any]:
    scores = result.get("category_scores", {})
    normalized = {}

    for key in WEIGHTS:
        try:
            value = int(round(float(scores.get(key, 0))))
        except (TypeError, ValueError):
            value = 0
        normalized[key] = max(0, min(100, value))

    result["category_scores"] = normalized
    result["ats_score"] = round(
        sum(normalized[k] * WEIGHTS[k] for k in WEIGHTS) / 100
    )

    for key in ("strengths", "improvements", "bullet_rewrites", "ats_checklist"):
        if not isinstance(result.get(key), list):
            result[key] = []

    if not isinstance(result.get("category_evidence"), dict):
        result["category_evidence"] = {}

    if not isinstance(result.get("keyword_analysis"), dict):
        result["keyword_analysis"] = {
            "present": [],
            "missing_or_weak": [],
            "notes": "",
        }

    return result


def analyze_resume(resume: str, job_description: str) -> dict[str, Any]:
    api_key = get_api_key()
    if not api_key:
        raise RuntimeError(
            "GEMINI_API_KEY is missing. Add it to Streamlit Secrets or your "
            "local environment."
        )

    client = genai.Client(api_key=api_key)

    response = client.models.generate_content(
        model=MODEL_NAME,
        contents=build_prompt(resume, job_description),
        config=types.GenerateContentConfig(
            temperature=0.2,
            response_mime_type="application/json",
        ),
    )

    text = (response.text or "").strip()
    text = re.sub(r"^```json\s*", "", text, flags=re.I)
    text = re.sub(r"\s*```$", "", text)

    if not text:
        raise RuntimeError("Gemini returned an empty response.")

    try:
        result = json.loads(text)
    except json.JSONDecodeError as exc:
        raise RuntimeError(
            "Gemini returned invalid JSON. Please try the analysis again."
        ) from exc

    if not isinstance(result, dict):
        raise RuntimeError("Unexpected response format from Gemini.")

    return normalize_result(result)


def render_results(result: dict[str, Any]) -> None:
    score = result["ats_score"]

    st.divider()
    st.subheader("🎯 ATS Readiness Score")
    st.metric("Estimated ATS Score", f"{score}/100")
    st.progress(score / 100)

    if score >= 85:
        st.success("Excellent ATS readiness.")
    elif score >= 70:
        st.info("Good ATS readiness, with some improvements recommended.")
    elif score >= 55:
        st.warning("The resume needs several improvements.")
    else:
        st.error("The resume has important ATS-readiness issues.")

    st.caption(
        "This is an AI-based ATS-readiness estimate, not a score from a specific "
        "employer's ATS. Different ATS platforms and job descriptions can produce "
        "different results."
    )

    st.subheader("📊 Score Breakdown")
    keys = list(WEIGHTS.keys())
    for start in range(0, len(keys), 4):
        cols = st.columns(4)
        for col, key in zip(cols, keys[start:start + 4]):
            with col:
                st.metric(
                    key.replace("_", " ").title(),
                    f"{result['category_scores'][key]}/100",
                )

    with st.expander("Why did I get these scores?"):
        evidence = result.get("category_evidence", {})
        for key in WEIGHTS:
            st.markdown(
                f"**{key.replace('_', ' ').title()} — "
                f"{result['category_scores'][key]}/100**"
            )
            st.write(evidence.get(key, "No explanation provided."))

    st.subheader("📝 Overall Assessment")
    st.write(result.get("summary", ""))

    left, right = st.columns(2)

    with left:
        st.markdown("### ✅ Strengths")
        for strength in result.get("strengths", []):
            st.success(strength)

    with right:
        st.markdown("### 🔧 Priority Improvements")
        for item in result.get("improvements", []):
            if isinstance(item, dict):
                st.markdown(
                    f"**{item.get('priority', 'Medium')}: "
                    f"{item.get('issue', '')}**"
                )
                st.write(item.get("recommendation", ""))
            else:
                st.write(item)

    st.subheader("🔑 Keyword Analysis")
    keywords = result.get("keyword_analysis", {})
    c1, c2 = st.columns(2)

    with c1:
        st.markdown("**Present / clearly evidenced**")
        present = keywords.get("present", [])
        st.write(", ".join(present) if present else "None identified.")

    with c2:
        st.markdown("**Missing or weak**")
        missing = keywords.get("missing_or_weak", [])
        st.write(", ".join(missing) if missing else "None identified.")

    if keywords.get("notes"):
        st.caption(keywords["notes"])

    st.subheader("✍️ Bullet Point Improvements")
    rewrites = result.get("bullet_rewrites", [])

    if not rewrites:
        st.write("No bullet rewrites were suggested.")
    else:
        for i, item in enumerate(rewrites, 1):
            if not isinstance(item, dict):
                continue
            with st.expander(f"Rewrite {i}"):
                st.markdown("**Original**")
                st.write(item.get("original", ""))
                st.markdown("**Improved**")
                st.write(item.get("improved", ""))
                st.caption(item.get("reason", ""))

    st.subheader("☑️ ATS Checklist")
    for item in result.get("ats_checklist", []):
        st.write(f"☐ {item}")


st.title("📄 Resume ATS Analyzer")
st.write(
    "Upload your resume and optionally add a job description to get an "
    "AI-powered ATS-readiness score and actionable improvements."
)

with st.sidebar:
    st.header("How it works")
    st.markdown(
        """
**1. Upload resume**

PDF, DOCX, TXT or MD.

**2. Add job description**

Optional, but recommended for keyword matching.

**3. Analyze**

Gemini reviews the resume.

**4. Improve**

Get a score, weaknesses, keywords, and bullet rewrites.
"""
    )
    st.divider()
    st.caption(f"AI model: {MODEL_NAME}")
    st.caption("No resume is permanently stored by this app.")

resume_file = st.file_uploader(
    "📎 Upload your resume",
    type=["pdf", "docx", "txt", "md"],
)

job_description = st.text_area(
    "💼 Job Description (Optional)",
    height=220,
    placeholder="Paste the job description here for job-specific ATS analysis...",
)

analyze = st.button(
    "🚀 Analyze Resume",
    type="primary",
    disabled=resume_file is None,
    use_container_width=True,
)

if analyze and resume_file:
    try:
        with st.spinner("Extracting resume text..."):
            resume_text = extract_text(resume_file)

        if len(resume_text) < 100:
            st.warning(
                "Only a small amount of text was extracted. "
                "Make sure your resume is readable and not image-only."
            )

        with st.spinner("Gemini is analyzing your resume..."):
            result = analyze_resume(resume_text, job_description)

        st.session_state["result"] = result
        st.session_state["filename"] = resume_file.name

    except Exception as exc:
        st.error(str(exc))

if "result" in st.session_state:
    st.caption(f"Analyzed: {st.session_state.get('filename', 'resume')}")
    render_results(st.session_state["result"])
